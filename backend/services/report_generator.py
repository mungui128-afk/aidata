import io
import os
import tempfile
from datetime import datetime
from typing import Any

FONT_REGISTERED = False


def _get_plt():
    import matplotlib

    matplotlib.use("Agg")
    import koreanize_matplotlib  # noqa: F401
    import matplotlib.pyplot as plt

    return plt


def _register_korean_font():
    global FONT_REGISTERED
    if FONT_REGISTERED:
        return "KoreanFont"

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_paths = [
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/malgunbd.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in font_paths:
        if os.path.exists(path):
            pdfmetrics.registerFont(TTFont("KoreanFont", path))
            FONT_REGISTERED = True
            return "KoreanFont"
    FONT_REGISTERED = True
    return "Helvetica"


def _format_date_kr(dt: datetime | None = None) -> str:
    d = dt or datetime.now()
    return f"{d.year}년 {d.month:02d}월 {d.day:02d}일"


def _format_currency(value: float) -> str:
    return f"{value:,.0f}원"


def create_chart_images(dashboard: dict[str, Any]) -> dict[str, bytes]:
    plt = _get_plt()
    charts: dict[str, bytes] = {}

    if dashboard.get("monthly_trend"):
        fig, ax = plt.subplots(figsize=(8, 4))
        months = [d["month"] for d in dashboard["monthly_trend"]]
        revenues = [d.get("revenue", 0) for d in dashboard["monthly_trend"]]
        ax.bar(months, revenues, color="#3b82f6", alpha=0.85)
        ax.set_title("월별 매출 추이", fontsize=14, fontweight="bold")
        ax.set_ylabel("매출 (원)")
        ax.tick_params(axis="x", rotation=45)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        charts["monthly"] = buf.getvalue()

    if dashboard.get("category_breakdown"):
        fig, ax = plt.subplots(figsize=(7, 4))
        labels = [d["name"] for d in dashboard["category_breakdown"][:8]]
        values = [d["value"] for d in dashboard["category_breakdown"][:8]]
        colors_list = plt.cm.Set3(range(len(labels)))
        ax.pie(values, labels=labels, autopct="%1.1f%%", colors=colors_list, startangle=90)
        ax.set_title("카테고리별 매출 비중", fontsize=14, fontweight="bold")
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        charts["category"] = buf.getvalue()

    if dashboard.get("product_top"):
        fig, ax = plt.subplots(figsize=(8, 4))
        products = [d["name"][:15] for d in dashboard["product_top"][:8]]
        values = [d["value"] for d in dashboard["product_top"][:8]]
        ax.barh(products[::-1], values[::-1], color="#10b981", alpha=0.85)
        ax.set_title("제품별 매출 TOP", fontsize=14, fontweight="bold")
        ax.set_xlabel("매출 (원)")
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        charts["product"] = buf.getvalue()

    return charts


def generate_docx(dashboard: dict[str, Any], report: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    title = doc.add_heading(report.get("title", "ERP 경영 분석 보고서"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"작성일: {_format_date_kr()}")
    doc.add_paragraph(f"분석 모델: {report.get('generated_by', 'Gemini AI')}")
    doc.add_paragraph("")

    doc.add_heading("경영 KPI 요약", level=1)
    kpis = dashboard["kpis"]
    kpi_table = doc.add_table(rows=1, cols=2)
    kpi_table.style = "Table Grid"
    hdr = kpi_table.rows[0].cells
    hdr[0].text = "지표"
    hdr[1].text = "값"
    kpi_rows = [
        ("총 매출", _format_currency(kpis["total_revenue"])),
        ("총 원가", _format_currency(kpis["total_cost"])),
        ("총 이익", _format_currency(kpis["total_profit"])),
        ("이익률", f"{kpis['profit_margin']}%"),
        ("거래 건수", f"{kpis['row_count']:,}건"),
    ]
    for label, value in kpi_rows:
        row = kpi_table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph("")
    doc.add_heading("경영진 요약", level=1)
    doc.add_paragraph(report.get("executive_summary", ""))

    charts = create_chart_images(dashboard)
    if charts:
        doc.add_heading("시각화 분석", level=1)
        with tempfile.TemporaryDirectory() as tmpdir:
            for key, data in charts.items():
                path = os.path.join(tmpdir, f"{key}.png")
                with open(path, "wb") as f:
                    f.write(data)
                doc.add_picture(path, width=Inches(5.5))
                doc.add_paragraph("")

    if dashboard.get("summary_table"):
        doc.add_heading("집계 데이터표", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        headers = ["구분", "매출", "수량", "이익"]
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for item in dashboard["summary_table"][:12]:
            row = tbl.add_row().cells
            row[0].text = str(item["group"])
            row[1].text = _format_currency(item["revenue"])
            row[2].text = f"{item['quantity']:,.0f}"
            row[3].text = _format_currency(item["profit"])

    for section in report.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        doc.add_paragraph(section.get("content", ""))
        for insight in section.get("insights", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(insight)

    doc.add_heading("권고사항", level=1)
    for rec in report.get("recommendations", []):
        doc.add_paragraph(rec, style="List Number")

    if report.get("risk_factors"):
        doc.add_heading("리스크 요인", level=1)
        for risk in report["risk_factors"]:
            doc.add_paragraph(risk, style="List Bullet")

    doc.add_heading("향후 전망", level=1)
    doc.add_paragraph(report.get("outlook", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf(dashboard: dict[str, Any], report: dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _register_korean_font()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleKR", parent=styles["Title"], fontName=font_name, fontSize=18, spaceAfter=12)
    heading_style = ParagraphStyle("HeadingKR", parent=styles["Heading1"], fontName=font_name, fontSize=14, spaceAfter=8, textColor=colors.HexColor("#1e40af"))
    body_style = ParagraphStyle("BodyKR", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16, spaceAfter=6)
    bullet_style = ParagraphStyle("BulletKR", parent=body_style, leftIndent=12, bulletIndent=6)

    story = []
    story.append(Paragraph(report.get("title", "ERP 경영 분석 보고서"), title_style))
    story.append(Paragraph(f"작성일: {_format_date_kr()} | 모델: {report.get('generated_by', 'Gemini AI')}", body_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("경영 KPI 요약", heading_style))
    kpis = dashboard["kpis"]
    kpi_data = [
        ["지표", "값"],
        ["총 매출", _format_currency(kpis["total_revenue"])],
        ["총 원가", _format_currency(kpis["total_cost"])],
        ["총 이익", _format_currency(kpis["total_profit"])],
        ["이익률", f"{kpis['profit_margin']}%"],
        ["거래 건수", f"{kpis['row_count']:,}건"],
    ]
    kpi_table = Table(kpi_data, colWidths=[80 * mm, 80 * mm])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4ff")]),
        ("ALIGN", (1, 1), (1, -1), "RIGHT"),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 16))

    story.append(Paragraph("경영진 요약", heading_style))
    story.append(Paragraph(report.get("executive_summary", ""), body_style))
    story.append(Spacer(1, 12))

    charts = create_chart_images(dashboard)
    if charts:
        story.append(Paragraph("시각화 분석", heading_style))
        for data in charts.values():
            story.append(Image(io.BytesIO(data), width=160 * mm, height=80 * mm))
            story.append(Spacer(1, 8))

    if dashboard.get("summary_table"):
        story.append(Paragraph("집계 데이터표", heading_style))
        table_data = [["구분", "매출", "수량", "이익"]]
        for item in dashboard["summary_table"][:12]:
            table_data.append([
                str(item["group"])[:20],
                _format_currency(item["revenue"]),
                f"{item['quantity']:,.0f}",
                _format_currency(item["profit"]),
            ])
        data_table = Table(table_data, colWidths=[45 * mm, 40 * mm, 35 * mm, 40 * mm])
        data_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ]))
        story.append(data_table)
        story.append(Spacer(1, 12))

    for section in report.get("sections", []):
        story.append(Paragraph(section.get("heading", ""), heading_style))
        story.append(Paragraph(section.get("content", ""), body_style))
        for insight in section.get("insights", []):
            story.append(Paragraph(f"• {insight}", bullet_style))

    story.append(Paragraph("권고사항", heading_style))
    for i, rec in enumerate(report.get("recommendations", []), 1):
        story.append(Paragraph(f"{i}. {rec}", body_style))

    if report.get("risk_factors"):
        story.append(Paragraph("리스크 요인", heading_style))
        for risk in report["risk_factors"]:
            story.append(Paragraph(f"• {risk}", bullet_style))

    story.append(Paragraph("향후 전망", heading_style))
    story.append(Paragraph(report.get("outlook", ""), body_style))

    doc.build(story)
    return buf.getvalue()
